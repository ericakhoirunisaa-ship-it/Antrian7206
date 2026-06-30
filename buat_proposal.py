from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Default font ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ─────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────

def set_cell_shading(cell, fill_hex='D9D9D9'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def identity_row(doc, number, label_text, value_text=None):
    """Single-cell shaded row like the cover table."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, 'D9D9D9')
    p    = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run  = p.add_run(f'{number}.  {label_text}')
    run.bold      = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if value_text:
        run2 = p.add_run(f' {value_text}')
        run2.bold      = False
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
    doc.add_paragraph('')   # small gap

def identity_row_multiline(doc, number, label_text, lines):
    """Shaded row with label on first line then detail lines inside same cell."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, 'D9D9D9')
    # First paragraph = label
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f'{number}.  {label_text}')
    r.bold      = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    # Extra lines
    for line in lines:
        p2 = cell.add_paragraph(line)
        p2.paragraph_format.left_indent = Cm(1)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        for run in p2.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    doc.add_paragraph('')

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold           = True
    r.font.name      = 'Times New Roman'
    r.font.size      = Pt(12)
    r.font.underline = False

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold      = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold      = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(1.0 + level * 0.5)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def spacer(doc):
    p = doc.add_paragraph('')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ═══════════════════════════════════════════════
# HALAMAN 1 — IDENTITAS INOVASI
# ═══════════════════════════════════════════════

identity_row(doc, 1,
    'Judul Inovasi:',
    'ANTRI-CENGKEH (Antrian Terpadu Real-time Informatif — '
    'Cepat Efisien Nomor Gilirannya Kiosk Elektronik Harian)')

identity_row(doc, 2, 'Nama Instansi:', 'Badan Pusat Statistik')

identity_row(doc, 3,
    'Nama Organisasi Penyelenggara Pelayanan [OPP]:',
    'Badan Pusat Statistik Kabupaten Tolitoli')

identity_row(doc, 4, 'Tipe Instansi:', 'Kementerian/Lembaga')

identity_row(doc, 5, 'Waktu mulai implementasi:', '1 Juli 2025')

identity_row_multiline(doc, 6, 'Identitas Inovator:', [
    'Nama  : Erika Khoirunisa',
    'NIP     : [diisi sesuai NIP]',
])

identity_row(doc, 7, 'Kelompok Inovasi:', 'Kelompok Umum')

identity_row_multiline(doc, 8, 'Kategori Inovasi:', [
    'Transformasi digital pelayanan publik',
])

identity_row_multiline(doc, 9, 'Target SDGs:', [
    'Industri, inovasi dan infrastruktur',
])

identity_row_multiline(doc, 10, 'Target Asta Cita:', [
    'Memperkuat pembangunan sumber daya manusia, sains, teknologi, pendidikan',
])

identity_row(doc, 11, 'Jenis Inovasi:', 'Inovasi digital')

identity_row(doc, 12, 'Sektor:', 'Pemerintahan umum/statistik')

identity_row(doc, 13, 'Link Video:', '[diisi setelah video dibuat]')

identity_row(doc, 14, 'Lokasi Inovasi:', '1.044565,120.822398')

doc.add_page_break()


# ═══════════════════════════════════════════════
# ASPEK 1: KEBARUAN
# ═══════════════════════════════════════════════

heading1(doc, 'ASPEK 1: KEBARUAN')

heading2(doc, '1.1 LATAR BELAKANG')

body(doc,
    'Pelayanan antrian di Pelayanan Statistik Terpadu (PST) BPS Kabupaten Tolitoli '
    'sebelumnya dilakukan secara manual menggunakan daftar tunggu berbasis kertas. '
    'Pengunjung yang datang tidak memiliki kepastian waktu giliran, sehingga menciptakan '
    'penumpukan antrean, ketidaknyamanan ruang tunggu, dan ketidakefisienan pelayanan. '
    'Petugas pun kesulitan memantau jumlah pengunjung secara real-time serta '
    'menghasilkan laporan kunjungan harian yang akurat dan terstruktur.')

body(doc,
    'ANTRI-CENGKEH hadir sebagai solusi digitalisasi sistem antrian yang terintegrasi '
    'dengan identitas daerah Kabupaten Tolitoli sebagai Kota Cengkeh. Sistem ini '
    'memungkinkan pengunjung mengambil nomor antrian secara mandiri melalui kiosk '
    'digital, petugas memantau dan memanggil antrian melalui dashboard, serta nomor '
    'antrian ditampilkan secara real-time pada layar digital di ruang tunggu. '
    'Data antrian tersimpan otomatis dan dapat digunakan untuk evaluasi serta pelaporan '
    'harian. Inovasi ini mendukung prinsip pelayanan prima, akuntabilitas, dan '
    'efisiensi kinerja harian BPS Kabupaten Tolitoli.')

heading2(doc, '1.2 TUJUAN')

body(doc, 'Tujuan :')
bullet(doc, 'Meningkatkan efisiensi dan kenyamanan proses antrian layanan PST.')
bullet(doc, 'Menyediakan basis data kunjungan yang terdokumentasi dan terstruktur secara digital.')
bullet(doc, 'Mendukung monitoring dan evaluasi pelayanan berbasis data aktual.')

spacer(doc)
body(doc, 'Target terukur :')
bullet(doc, '100% antrian pengunjung PST tercatat secara digital.')
bullet(doc, 'Pengurangan waktu tunggu rata-rata pengunjung hingga 40%.')
bullet(doc, 'Laporan kunjungan harian dapat dihasilkan dalam waktu kurang dari 1 menit.')
bullet(doc, 'Kepuasan pengunjung terhadap sistem antrian meningkat minimal 20%.')

heading2(doc, '1.3 CARA IMPLEMENTASI')

bullet(doc, 'Sistem berbasis aplikasi web (HTML/CSS/JavaScript) yang berjalan pada '
            'jaringan lokal kantor tanpa memerlukan instalasi perangkat lunak khusus.')
bullet(doc, 'Pengunjung memilih jenis layanan (Konsultasi Statistik, Perpustakaan '
            'Statistik, Rekomendasi Statistik, atau Pelayanan Lainnya) dan mengambil '
            'nomor antrian secara mandiri melalui kiosk digital berupa layar sentuh atau tablet.')
bullet(doc, 'Nomor antrian yang aktif ditampilkan secara real-time pada layar antrian '
            'digital yang dipasang di ruang tunggu.')
bullet(doc, 'Petugas memanggil dan mengelola antrian melalui dashboard petugas yang '
            'dapat diakses di komputer meja layanan.')
bullet(doc, 'Data seluruh antrian tersimpan otomatis dan dapat diekspor dalam format '
            'CSV untuk keperluan evaluasi dan pelaporan.')
bullet(doc, 'Admin dapat memantau seluruh aktivitas antrian, mengelola akun, dan '
            'melakukan reset data harian melalui panel administrator.')
bullet(doc, 'Review bulanan dilakukan untuk perbaikan fitur atau penyesuaian SOP layanan.')

heading2(doc, '1.4 KEUNGGULAN IDE/GAGASAN')

bullet(doc, 'Gratis dan hemat biaya — dikembangkan menggunakan teknologi web terbuka '
            'tanpa biaya lisensi perangkat lunak tambahan.')
bullet(doc, 'Mandiri dari internet — berjalan sepenuhnya pada jaringan lokal kantor '
            'sehingga tidak terganggu oleh koneksi internet yang tidak stabil.')
bullet(doc, 'Real-time monitoring antrian oleh petugas dan pimpinan melalui dashboard.')
bullet(doc, 'Terintegrasi dengan sistem evaluasi pelayanan melalui ekspor laporan CSV otomatis.')
bullet(doc, 'Mudah direplikasi ke seluruh unit kerja BPS tanpa infrastruktur yang rumit.')
bullet(doc, 'Antarmuka ramah pengguna — dirancang sederhana dan intuitif untuk segala '
            'usia dan latar belakang pendidikan pengunjung.')

doc.add_page_break()


# ═══════════════════════════════════════════════
# ASPEK 2: EFEKTIFITAS DAN MANFAAT
# ═══════════════════════════════════════════════

heading1(doc, 'ASPEK 2: EFEKTIFITAS DAN MANFAAT')

heading2(doc, '2.1 INDIKATOR MONEV')

body(doc, '   Instrumen :')
bullet(doc, 'Monitoring harian oleh petugas PST melalui dashboard antrian digital.', level=1)
bullet(doc, 'Evaluasi mingguan oleh admin sistem terhadap laporan CSV antrian.', level=1)
bullet(doc, 'Evaluasi bulanan oleh pimpinan OPP terhadap rekap data kunjungan.', level=1)

body(doc, '   Indikator :')
bullet(doc, 'Jumlah antrian yang tercatat per hari dan per jenis layanan.', level=1)
bullet(doc, 'Waktu rata-rata tunggu pengunjung dari pengambilan nomor hingga dipanggil.', level=1)
bullet(doc, 'Tingkat pemanfaatan laporan antrian digital untuk pengambilan keputusan pelayanan.', level=1)

heading2(doc, '2.2 DAMPAK INOVASI')

body(doc, 'a. Bentuk Dampak :', bold=True)
bullet(doc, 'Meningkatkan kecepatan, keteraturan, dan kenyamanan proses antrian pengunjung PST.')
bullet(doc, 'Mendukung pengambilan keputusan manajerial berbasis data kunjungan yang akurat.')

spacer(doc)
body(doc, 'b. Capaian Output dan Outcome :', bold=True)

body(doc, 'Sebelum ANTRI-CENGKEH :', bold=True)
bullet(doc, 'Antrian manual menggunakan daftar tunggu kertas.')
bullet(doc, 'Data pengunjung tidak terstruktur, rentan hilang atau tidak lengkap.')
bullet(doc, 'Petugas kesulitan memantau jumlah antrian aktif secara bersamaan.')
bullet(doc, 'Tidak ada laporan kunjungan harian yang tersedia secara otomatis.')

body(doc, 'Sesudah ANTRI-CENGKEH :', bold=True)
bullet(doc, 'Pengambilan nomor antrian digital secara mandiri dan cepat oleh pengunjung.')
bullet(doc, 'Data kunjungan tersimpan otomatis, terstruktur, dan siap dianalisis.')
bullet(doc, 'Petugas memantau seluruh antrian real-time dari satu dashboard.')
bullet(doc, 'Laporan kunjungan harian tersedia otomatis dan dapat diekspor setiap saat.')

spacer(doc)
body(doc, 'c. Dampak terhadap Asta Cita :', bold=True)
bullet(doc, 'Mendukung reformasi birokrasi melalui digitalisasi dan efisiensi layanan publik BPS.')
bullet(doc, 'Mendorong peningkatan kapasitas SDM dalam pemanfaatan teknologi informasi '
            'untuk pelayanan prima berbasis data.')

doc.add_page_break()


# ═══════════════════════════════════════════════
# ASPEK 3: ADAPTABILITAS
# ═══════════════════════════════════════════════

heading1(doc, 'ASPEK 3: ADAPTABILITAS')

heading2(doc, '3.1 DIFUSI DAN REPLIKASI')

body(doc, '         a. Potensi Replikasi :')
bullet(doc, 'Tinggi; dapat digunakan oleh seluruh unit kerja BPS provinsi dan '
            'kabupaten/kota di Indonesia tanpa biaya pengembangan besar.', level=1)
bullet(doc, 'Dapat direplikasi oleh instansi pemerintah lain yang memiliki kebutuhan '
            'sistem antrian pelayanan publik.', level=1)

body(doc, '         b. Upaya Difusi :')
bullet(doc, 'Sosialisasi melalui grup WhatsApp internal BPS dan presentasi di forum '
            'peningkatan kapasitas satuan kerja.', level=1)
bullet(doc, 'Transfer pengetahuan melalui pelatihan singkat dan penyediaan template '
            'sistem yang siap pakai.', level=1)

body(doc, '         c. Jumlah unit kerja dan/atau instansi yang telah mereplikasi inovasi :')
body(doc, '         Belum ada unit kerja dan/atau instansi yang mereplikasi.')

doc.add_paragraph('')


# ═══════════════════════════════════════════════
# ASPEK 4: KEBERLANJUTAN
# ═══════════════════════════════════════════════

heading1(doc, 'ASPEK 4: KEBERLANJUTAN')

heading2(doc, '4.1 SUMBER DAYA')

body(doc, 'Penggunaan sumber daya :')
bullet(doc, 'Sarpras: komputer/laptop, layar monitor tambahan untuk display antrian, '
            'perangkat tablet/layar sentuh untuk kiosk pengambilan nomor, '
            'koneksi jaringan lokal (LAN/WiFi) kantor.')
bullet(doc, 'SDI: petugas operator PST dan admin sistem antrian.')
bullet(doc, 'SDA: biaya operasional sangat rendah — sistem berbasis teknologi web terbuka '
            'tanpa biaya lisensi perangkat lunak.')
bullet(doc, 'Informasi: data antrian tersimpan di perangkat lokal dan dapat dicadangkan '
            'secara berkala dalam format CSV.')

heading2(doc, '4.2 STRATEGI KEBERLANJUTAN')

body(doc, '      1) Strategi institusional :')
bullet(doc, 'Pencantuman penggunaan ANTRI-CENGKEH dalam SOP Pelayanan PST '
            'BPS Kabupaten Tolitoli.', level=1)
bullet(doc, 'Dukungan regulasi internal sebagai standar baku sistem antrian '
            'pelayanan tamu.', level=1)

body(doc, '      2) Strategi manajerial :')
bullet(doc, 'Pelatihan SDM secara berkala dalam pengoperasian sistem dan '
            'pengelolaan data antrian.', level=1)
bullet(doc, 'SOP pengelolaan, pemantauan harian, dan backup data antrian secara berkala.', level=1)

body(doc, '      3) Strategi sosial :')
bullet(doc, 'Sosialisasi kepada seluruh pengunjung PST tentang tata cara '
            'penggunaan kiosk antrian digital.', level=1)
bullet(doc, 'Terbuka untuk replikasi dan kolaborasi lintas instansi dalam rangka '
            'peningkatan pelayanan publik berbasis teknologi.', level=1)


# ── Save ──
output_path = r'C:\Users\ERIKA A K\Documents\7206\Sistem Antrian\Proposal Inovasi ANTRI-CENGKEH.docx'
doc.save(output_path)
print(f'Dokumen berhasil dibuat: {output_path}')
