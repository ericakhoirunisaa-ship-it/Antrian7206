# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE   = r'C:\Users\ERIKA A K\Documents\7206\Sistem Antrian'
JUDUL  = 'ANTRI-CENGKEH (Antrian Terpadu Real-time Informatif — Cepat Efisien Nomor Gilirannya Kiosk Elektronik Harian)'
INSTANSI = 'Badan Pusat Statistik Kabupaten Tolitoli'

# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(3.0)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(4.0)
        s.right_margin  = Cm(3.0)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    return doc

def add_hline(doc):
    """Bottom border on an empty paragraph = horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def para(doc, text='', bold=False, italic=False, align='justify',
         size=12, sb=0, sa=6, indent=False, color=None):
    p = doc.add_paragraph()
    if   align == 'center':  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent: p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
    return p

def mixed_para(doc, parts, align='justify', sb=4, sa=4, indent=False):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    if   align == 'center':  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent: p.paragraph_format.first_line_indent = Cm(1.25)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def heading(doc, text, sb=10, sa=4):
    return para(doc, text, bold=True, align='left', sb=sb, sa=sa)

def bullet(doc, text, level=1, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(level * 0.75)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    prefix = '•  '
    if bold_prefix:
        r0 = p.add_run(prefix + bold_prefix)
        r0.bold = True; r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(text)
        r1.bold = False; r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    else:
        r = p.add_run(prefix + text)
        r.font.name='Times New Roman'; r.font.size=Pt(12)

def sub_bullet(doc, text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(level * 0.75)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('o  ' + text)
    r.font.name='Times New Roman'; r.font.size=Pt(12)

def numbered(doc, num, text, bold_prefix=None, level=1, sb=2, sa=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent        = Cm(level * 0.75)
    p.paragraph_format.first_line_indent  = Cm(-0.5)
    p.paragraph_format.space_before       = Pt(sb)
    p.paragraph_format.space_after        = Pt(sa)
    prefix = f'{num}.  '
    if bold_prefix:
        r0 = p.add_run(prefix + bold_prefix)
        r0.bold=True; r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(text)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    else:
        r = p.add_run(prefix + text)
        r.font.name='Times New Roman'; r.font.size=Pt(12)

def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(11)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(val, tuple):   # (text, bold)
                r = p.add_run(val[0]); r.bold=val[1]
            else:
                r = p.add_run(val)
            r.font.name='Times New Roman'; r.font.size=Pt(11)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def doc_header(doc, title_line):
    para(doc, title_line, bold=True, align='left', sb=0, sa=4)
    mixed_para(doc, [('Judul Inovasi: ', True, False), (JUDUL, False, False)],
               align='left', sb=0, sa=4)
    mixed_para(doc, [('Instansi Pengusul: ', True, False), (INSTANSI, False, False)],
               align='left', sb=0, sa=4)
    add_hline(doc)

def save(doc, filename):
    path = os.path.join(BASE, filename)
    doc.save(path)
    print(f'  OK  {filename}')


# ══════════════════════════════════════════════════════════════
# 1. KEUNGGULAN IDE
# ══════════════════════════════════════════════════════════════
def buat_keunggulan():
    doc = new_doc()
    doc_header(doc, 'BUKTI DUKUNG KEUNGGULAN IDE / GAGASAN INOVASI')

    heading(doc, 'I. KEUNGGULAN IDE YANG DITAWARKAN')
    para(doc,
         'ANTRI-CENGKEH menawarkan pendekatan baru yang inovatif dan efisien dalam '
         'pengelolaan antrian layanan di lingkungan Pelayanan Statistik Terpadu (PST) '
         'BPS Kabupaten Tolitoli. Ide ini membawa kebaruan dalam beberapa aspek berikut:',
         align='justify', indent=True, sa=6)

    items = [
        ('Digitalisasi Antrian Manual Tanpa Biaya Tinggi',
         [('Menggantikan sistem antrian kertas dengan kiosk digital berbasis web yang '
           'dapat berjalan di perangkat yang sudah tersedia di kantor.'),
          ('Tidak memerlukan infrastruktur mahal, server khusus, atau koneksi internet.')]),
        ('Penomoran Antrian Otomatis per Jenis Layanan',
         [('Nomor antrian diberikan secara otomatis dan terpisah per jenis layanan '
           '(Konsultasi, Perpustakaan, Rekomendasi Statistik, dan Lainnya).'),
          ('Menghilangkan risiko penomoran ganda atau antrian yang terlewat.')]),
        ('Tampilan Antrian Real-Time di Layar Digital',
         [('Nomor yang sedang dilayani tampil secara real-time pada layar antrian '
           'di ruang tunggu tanpa perlu di-refresh manual.'),
          ('Disertai pengumuman suara otomatis saat nomor dipanggil.')]),
        ('Dashboard Petugas Terintegrasi',
         [('Petugas mengelola seluruh antrian dari satu dashboard tanpa perlu '
           'keluar ruangan atau membuat catatan manual.'),
          ('Data antrian tersedia langsung untuk laporan dan evaluasi harian.')]),
        ('Praktis, Gratis, dan Ramah Pengguna',
         [('Pengambilan nomor antrian hanya membutuhkan beberapa detik.'),
          ('Antarmuka dirancang sederhana dan intuitif untuk semua kalangan '
           'pengunjung tanpa perlu keahlian teknis.')]),
    ]
    for i, (title, subs) in enumerate(items, 1):
        numbered(doc, i, '', bold_prefix=title, sb=4, sa=2)
        for s in subs: sub_bullet(doc, s, level=2)

    add_hline(doc)
    heading(doc, 'II. PERBEDAAN DENGAN SISTEM SEBELUMNYA')
    add_table(doc,
        ['Aspek', 'Sistem Lama (Manual)', 'ANTRI-CENGKEH (Digital)'],
        [('Pengambilan Nomor',  'Tulis nomor di kertas secara manual',        'Otomatis via kiosk digital'),
         ('Tampilan Antrian',   'Dipanggil secara verbal oleh petugas',       'Layar digital + suara otomatis'),
         ('Pencatatan Kunjungan','Tidak ada pencatatan terstruktur',           'Tersimpan otomatis di sistem'),
         ('Laporan Harian',     'Tidak tersedia',                             'Ekspor CSV otomatis kapan saja'),
         ('Monitoring Petugas', 'Tidak ada',                                  'Dashboard real-time terintegrasi'),
         ('Biaya Operasional',  'Kertas dan alat tulis',                      'Nol biaya tambahan (web gratis)'),
        ], col_widths=[4.5, 5.5, 6.0])

    add_hline(doc)
    heading(doc, 'III. POTENSI PENGEMBANGAN GAGASAN')
    para(doc, 'ANTRI-CENGKEH memiliki potensi dikembangkan ke:', align='left', sa=4)
    for t in ['Integrasi dengan sistem layanan statistik daring BPS Pusat.',
              'Penambahan fitur survei kepuasan pengunjung secara digital langsung setelah dilayani.',
              'Notifikasi antrian melalui pesan pendek (SMS) atau WhatsApp.',
              'Integrasi dengan sistem presensi dan evaluasi kinerja pelayanan harian.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'IV. KEBARUAN INOVASI DALAM KONTEKS INSTANSI')
    for t in ['Belum ada sistem antrian digital di BPS Kabupaten Tolitoli sebelumnya.',
              'Merupakan sistem pertama yang mengintegrasikan kiosk antrian, layar display real-time, '
              'dashboard petugas, dan laporan dalam satu platform terpadu berbasis web lokal.',
              'Inovasi muncul dari kebutuhan lokal namun solusinya dapat diterapkan secara luas '
              'di seluruh unit kerja BPS maupun instansi pelayanan publik lainnya.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'PENUTUP', sb=6)
    para(doc,
         'ANTRI-CENGKEH membuktikan bahwa transformasi digital pelayanan publik dapat dilakukan '
         'secara cepat, murah, dan berdampak besar. Sistem ini menjadi contoh inovasi yang dapat '
         'ditiru oleh satuan kerja BPS lain maupun instansi pelayanan publik lainnya.',
         align='justify', indent=True)

    add_hline(doc)
    para(doc, 'Lampiran :', italic=True, align='left', sb=4, sa=2)
    for t in ['Tabel perbandingan sistem antrian manual vs digital.',
              'Video demonstrasi penggunaan sistem ANTRI-CENGKEH.',
              'Tangkapan layar kiosk antrian, layar display, dan dashboard petugas.']:
        bullet(doc, t)

    save(doc, 'BD 1 - KEUNGGULAN IDE.docx')


# ══════════════════════════════════════════════════════════════
# 2. TUJUAN INOVASI
# ══════════════════════════════════════════════════════════════
def buat_tujuan():
    doc = new_doc()
    doc_header(doc, 'BUKTI DUKUNG TUJUAN INOVASI')

    heading(doc, 'I. TUJUAN UMUM INOVASI')
    para(doc,
         'Inovasi ANTRI-CENGKEH bertujuan untuk mentransformasi sistem antrian pelayanan '
         'statistik dari manual menjadi digital agar lebih efisien, tertib, akuntabel, '
         'dan mendukung evaluasi pelayanan berbasis data secara harian.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, 'II. TUJUAN KHUSUS INOVASI')
    items = [
        ('Meningkatkan efisiensi dan kenyamanan proses antrian layanan PST.',
         'Mengurangi waktu tunggu dan ketidakpastian giliran bagi pengunjung.'),
        ('Menyediakan basis data kunjungan yang tertata dan terdokumentasi secara digital.',
         'Memastikan seluruh data antrian tersimpan dalam sistem yang terstruktur dan dapat diakses kapan saja.'),
        ('Mendukung pengambilan keputusan berbasis data secara harian.',
         'Melalui dashboard dan laporan CSV yang ter-update otomatis dari data antrian.'),
        ('Menjamin akuntabilitas dan transparansi pelayanan pengunjung PST.',
         'Memberikan akses monitoring real-time bagi pimpinan dan petugas.'),
    ]
    for i, (title, sub) in enumerate(items, 1):
        numbered(doc, i, '', bold_prefix=title, sb=4, sa=2)
        sub_bullet(doc, sub, level=2)

    add_hline(doc)
    heading(doc, 'III. OUTCOME YANG DIHARAPKAN')
    for i, t in enumerate([
        'Pelayanan antrian yang lebih cepat, tertib, dan profesional.',
        'Evaluasi harian kinerja petugas PST dapat dilakukan dengan data antrian nyata.',
        'Terciptanya sistem monitoring antrian yang dapat diakses pimpinan setiap saat.',
    ], 1):
        numbered(doc, i, t)

    add_hline(doc)
    heading(doc, 'IV. OUTPUT YANG DIHASILKAN')
    items4 = [
        ('Kiosk digital pengambilan nomor antrian ', 'yang terintegrasi dengan sistem.'),
        ('Layar antrian digital ', 'yang menampilkan nomor secara real-time di ruang tunggu.'),
        ('Dashboard petugas ', 'untuk mengelola dan memanggil antrian layanan PST.'),
        ('Laporan antrian harian otomatis ', 'dalam format CSV yang siap diunduh kapan saja.'),
    ]
    for i, (bold_t, norm_t) in enumerate(items4, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before      = Pt(2)
        p.paragraph_format.space_after       = Pt(2)
        r0 = p.add_run(f'{i}.  {bold_t}'); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(norm_t)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'V. TARGET TERUKUR')
    add_table(doc,
        ['Indikator', 'Target'],
        [('Digitalisasi antrian',         '100% antrian pengunjung PST tercatat secara digital'),
         ('Efisiensi waktu tunggu',        'Waktu tunggu rata-rata pengunjung berkurang hingga 40%'),
         ('Kecepatan laporan evaluasi',    'Laporan dapat dihasilkan dalam waktu < 1 menit'),
         ('Keakuratan data antrian',       '> 95% antrian tercatat lengkap dan akurat'),
        ], col_widths=[6.0, 10.0])

    add_hline(doc)
    heading(doc, 'VI. KAITAN DENGAN VISI MISI INSTANSI & TARGET ASTA CITA')
    para(doc,
         'Inovasi ini sejalan dengan misi reformasi birokrasi dan transformasi digital '
         'pelayanan publik. Selain itu, inovasi mendukung pencapaian target ',
         align='justify', indent=True, sa=2)
    # Asta Cita in bold inline
    p = doc.paragraphs[-1]
    p.clear()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run('Inovasi ini sejalan dengan misi reformasi birokrasi dan transformasi digital '
                   'pelayanan publik. Selain itu, inovasi mendukung pencapaian target ')
    r0.font.name='Times New Roman'; r0.font.size=Pt(12)
    r1 = p.add_run('Asta Cita'); r1.bold=True
    r1.font.name='Times New Roman'; r1.font.size=Pt(12)
    r2 = p.add_run(', khususnya:')
    r2.font.name='Times New Roman'; r2.font.size=Pt(12)
    for t in ['Meningkatkan kualitas pelayanan publik berbasis teknologi informasi.',
              'Memperkuat kapasitas SDM dan infrastruktur teknologi statistik.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'PENUTUP', sb=6)
    para(doc,
         'Dokumen ini memberikan gambaran menyeluruh mengenai arah dan harapan dari '
         'implementasi inovasi ANTRI-CENGKEH. Dengan indikator terukur dan hasil nyata '
         'di lapangan, inovasi ini telah membawa perubahan signifikan dalam proses '
         'pelayanan antrian di lingkungan PST BPS Kabupaten Tolitoli.',
         align='justify', indent=True)

    save(doc, 'BD 2 - TUJUAN INOVASI.docx')


# ══════════════════════════════════════════════════════════════
# 3. CARA KERJA INOVASI
# ══════════════════════════════════════════════════════════════
def buat_cara_kerja():
    doc = new_doc()
    doc_header(doc, 'BUKTI DUKUNG CARA KERJA INOVASI')

    heading(doc, 'I. PENJELASAN UMUM CARA KERJA')
    para(doc,
         'ANTRI-CENGKEH adalah sistem antrian digital untuk Pelayanan Statistik Terpadu (PST) '
         'yang memanfaatkan teknologi web berbasis jaringan lokal kantor. Prosesnya melibatkan '
         'pengambilan nomor antrian secara mandiri oleh pengunjung, penampilan nomor secara '
         'real-time di layar antrian, dan pengelolaan antrian oleh petugas melalui dashboard '
         'terintegrasi, tanpa memerlukan koneksi internet maupun server tambahan.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, 'II. TAHAPAN IMPLEMENTASI SISTEM ANTRI-CENGKEH')

    tahapan = [
        ('Persiapan Sistem:', [
            'Seluruh file sistem (index.html, display.html, officer.html, admin.html, js/state.js) disiapkan di komputer kantor.',
            'Semua perangkat (kiosk/tablet, layar antrian, komputer petugas) dihubungkan ke jaringan WiFi/LAN lokal yang sama.',
            'Petugas membuka modul sesuai perangkat yang digunakan dan memastikan sistem berjalan normal.',
        ]),
        ('Pelaksanaan Antrian:', [
            'Pengunjung memilih jenis layanan dan mengambil nomor antrian secara mandiri melalui kiosk digital.',
            'Sistem secara otomatis menghasilkan nomor urut berdasarkan jenis layanan (K – Konsultasi, P – Perpustakaan, R – Rekomendasi, L – Lainnya).',
            'Tiket antrian digital ditampilkan di layar kiosk beserta informasi layanan dan estimasi antrian.',
        ]),
        ('Pemanggilan Antrian:', [
            'Petugas memantau daftar antrian dari dashboard dan memanggil pengunjung berikutnya per jenis layanan.',
            'Nomor yang dipanggil tampil secara otomatis di layar antrian ruang tunggu disertai pengumuman suara.',
            'Petugas dapat memanggil ulang, melewati, atau menandai antrian selesai dari dashboard.',
        ]),
        ('Penyelesaian dan Pelaporan:', [
            'Petugas menandai antrian sebagai selesai setelah pelayanan diberikan kepada pengunjung.',
            'Seluruh data antrian tersimpan otomatis dan dapat diekspor dalam format CSV oleh admin.',
            'Admin memantau laporan, statistik kunjungan, dan mereset data pada akhir hari kerja.',
        ]),
    ]
    for i, (title, subs) in enumerate(tahapan, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f'{i}. {title}'); r.bold=True
        r.font.name='Times New Roman'; r.font.size=Pt(12)
        for s in subs: bullet(doc, s, level=1)

    add_hline(doc)
    heading(doc, 'III. ILUSTRASI ALUR KERJA (Step-by-step)')
    steps = [
        'Pengunjung datang ke kantor BPS dan menuju kiosk antrian digital.',
        'Pengunjung memilih jenis layanan yang dibutuhkan di layar kiosk.',
        'Sistem mencetak nomor antrian digital secara otomatis sesuai jenis layanan.',
        'Pengunjung menunggu di ruang tunggu sambil memantau layar antrian.',
        'Layar antrian menampilkan nomor yang sedang dipanggil beserta pengumuman suara otomatis.',
        'Petugas melayani pengunjung dan menandai antrian selesai di dashboard.',
        'Data antrian tersimpan otomatis sebagai arsip dan laporan harian sistem.',
    ]
    for i, s in enumerate(steps, 1): numbered(doc, i, s)

    add_hline(doc)
    heading(doc, 'IV. PENGELOLAAN SISTEM')
    roles = [
        ('Admin Sistem',          ': Bertanggung jawab terhadap pengaturan, monitoring, ekspor data, dan reset antrian harian.'),
        ('Petugas PST',           ': Memanggil dan mengelola antrian melalui dashboard, memastikan pelayanan berjalan tertib.'),
        ('Koordinator Fungsional',': Mengecek dan menganalisis data antrian untuk evaluasi kinerja pelayanan harian.'),
    ]
    for role, desc in roles:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.75)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r0 = p.add_run('•  ' + role); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(desc)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'V. FLEKSIBILITAS DAN REPLIKASI')
    para(doc, 'Sistem ini dapat diterapkan dengan:', align='left', sa=4)
    for t in ['Infrastruktur minimal (hanya butuh perangkat komputer/tablet dan jaringan lokal kantor).',
              'Teknologi web terbuka tanpa biaya lisensi perangkat lunak apapun.',
              'File sistem dapat disalin dan langsung digunakan di unit kerja lain tanpa modifikasi rumit.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'PENUTUP', sb=6)
    para(doc,
         'Cara kerja inovasi ANTRI-CENGKEH telah dirancang untuk menyederhanakan proses '
         'antrian pelayanan statistik secara menyeluruh. Sistem ini mudah diterapkan, '
         'efisien, dan memiliki fleksibilitas tinggi untuk dikembangkan lebih lanjut.',
         align='justify', indent=True)

    add_hline(doc)
    para(doc, 'Lampiran :', italic=True, align='left', sb=4, sa=2)
    for t in ['Tangkapan layar kiosk antrian dan layar antrian digital.',
              'Panduan penggunaan sistem ANTRI-CENGKEH.',
              'Contoh tiket antrian digital.',
              'Contoh laporan CSV harian.']:
        bullet(doc, t)

    save(doc, 'BD 3 - CARA KERJA INOVASI.docx')


# ══════════════════════════════════════════════════════════════
# 4. LATAR BELAKANG INOVASI
# ══════════════════════════════════════════════════════════════
def buat_latar_belakang():
    doc = new_doc()
    doc_header(doc, 'BUKTI DUKUNG LATAR BELAKANG INOVASI')

    heading(doc, 'I. LATAR BELAKANG MASALAH')
    para(doc,
         'Sebelum adanya inovasi ANTRI-CENGKEH, proses antrian pengunjung di Pelayanan '
         'Statistik Terpadu (PST) BPS Kabupaten Tolitoli masih dilakukan secara manual '
         'menggunakan daftar tunggu berbasis kertas. Cara ini menghadirkan sejumlah permasalahan:',
         align='justify', indent=True, sa=4)

    masalah = [
        ('Inefisiensi dan ketidaktertiban antrian: ',
         'Pengunjung tidak memiliki kepastian waktu giliran sehingga menimbulkan '
         'penumpukan dan ketidaknyamanan di ruang tunggu PST.'),
        ('Tidak ada pencatatan kunjungan yang terstruktur: ',
         'Data pengunjung tidak tercatat secara sistematis sehingga sulit dilakukan '
         'evaluasi dan pelaporan pelayanan harian.'),
        ('Tidak mendukung evaluasi harian yang cepat: ',
         'Data tidak tersedia secara real-time, menyulitkan analisis dan pengambilan '
         'keputusan berbasis data.'),
        ('Kurangnya integrasi dan akuntabilitas: ',
         'Tidak tersedia sistem untuk memastikan ketertiban, kecepatan, serta '
         'kemudahan monitoring antrian oleh pimpinan.'),
    ]
    for i, (bold_t, norm_t) in enumerate(masalah, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before      = Pt(2)
        p.paragraph_format.space_after       = Pt(2)
        r0 = p.add_run(f'{i}.  {bold_t}'); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(norm_t)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    para(doc,
         'Masalah ini berdampak langsung terhadap kualitas pelayanan publik dan pengambilan '
         'keputusan internal, terutama dalam konteks evaluasi dan pengawasan harian pelayanan PST.',
         align='justify', indent=True, sb=6)

    add_hline(doc)
    heading(doc, 'II. PERUMUSAN SOLUSI MELALUI INOVASI')
    para(doc,
         'Inovasi ANTRI-CENGKEH dikembangkan sebagai solusi digital untuk mengatasi '
         'permasalahan tersebut. Sistem ini memanfaatkan teknologi web berbasis lokal untuk '
         'penomoran antrian otomatis, pengelolaan antrian real-time, dan pelaporan digital. '
         'ANTRI-CENGKEH tidak hanya sekadar digitalisasi antrian, tetapi juga berfungsi '
         'sebagai alat monitoring dan evaluasi pelayanan harian yang praktis dan efisien.',
         align='justify', indent=True, sa=4)
    para(doc, 'Langkah-langkah perumusan inovasi dilakukan melalui:', align='left', sa=4)
    solusi = [
        ('Identifikasi permasalahan di lapangan ', 'oleh tim pelayanan PST BPS Kabupaten Tolitoli.'),
        ('Riset solusi teknologi ringan dan gratis ', 'yang dapat diterapkan tanpa infrastruktur mahal.'),
        ('Pengembangan dan pengujian awal ', 'sistem pada perangkat yang tersedia di kantor.'),
        ('Penyusunan SOP dan panduan penggunaan sistem.', ''),
        ('Evaluasi internal dan perbaikan berkelanjutan.', ''),
    ]
    for i, (bold_t, norm_t) in enumerate(solusi, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before      = Pt(2)
        p.paragraph_format.space_after       = Pt(2)
        r0 = p.add_run(f'{i}.  {bold_t}'); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        if norm_t:
            r1 = p.add_run(norm_t)
            r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'III. IMPLEMENTASI DAN PERKEMBANGAN INOVASI')
    para(doc,
         'Sejak implementasi pada 1 Juli 2025, ANTRI-CENGKEH telah diterapkan secara '
         'menyeluruh untuk setiap kunjungan di kantor BPS Kabupaten Tolitoli. '
         'Implementasi dilakukan secara bertahap:',
         align='justify', indent=True, sa=4)
    tahapan = [
        ('Tahap 1 (Perencanaan): ',   'Pengembangan sistem, penyusunan SOP, dan pelatihan petugas PST.'),
        ('Tahap 2 (Uji Coba): ',      'Penggunaan sistem pada layanan PST selama 1 minggu uji coba.'),
        ('Tahap 3 (Implementasi Penuh): ', 'Seluruh antrian pengunjung PST wajib melalui sistem ANTRI-CENGKEH.'),
        ('Tahap 4 (Evaluasi dan Penyempurnaan): ',
         'Review berkala dilakukan untuk menyesuaikan fitur berdasarkan umpan balik petugas dan pengunjung.'),
    ]
    for i, (bold_t, norm_t) in enumerate(tahapan, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before      = Pt(2)
        p.paragraph_format.space_after       = Pt(2)
        r0 = p.add_run(f'{i}.  {bold_t}'); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(norm_t)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'IV. DAMPAK DAN MANFAAT')
    for t in ['100% antrian pengunjung PST tercatat secara digital.',
              'Waktu tunggu rata-rata pengunjung berkurang hingga 40%.',
              'Laporan antrian harian dapat diakses dalam waktu < 1 menit.',
              'Meningkatkan ketertiban dan kenyamanan ruang tunggu PST.',
              'Mempercepat proses evaluasi dan pelaporan pelayanan harian.']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.75)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run('•  ' + t); r.bold=True
        r.font.name='Times New Roman'; r.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'V. KESESUAIAN DENGAN TUJUAN PEMBANGUNAN NASIONAL DAN PRIORITAS INSTANSI')
    para(doc, 'ANTRI-CENGKEH mendukung:', align='left', sa=4)
    for t in ['Transformasi digital pelayanan publik.',
              'Reformasi birokrasi dan efisiensi kerja.',
              'Pemanfaatan data dalam pengambilan keputusan.',
              'Penguatan kapasitas SDM dan infrastruktur teknologi.',
              'Identitas kedaerahan Kabupaten Tolitoli sebagai Kota Cengkeh.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'VI. PENUTUP', sb=6)
    para(doc,
         'ANTRI-CENGKEH bukan hanya solusi digitalisasi antrian, tetapi menjadi instrumen '
         'penting dalam meningkatkan efisiensi layanan dan mendukung sistem evaluasi '
         'berbasis data. Dengan kebutuhan sumber daya yang minimal dan sistem yang mudah '
         'direplikasi, inovasi ini memiliki potensi besar untuk diadaptasi di unit kerja '
         'BPS lain maupun instansi pelayanan publik secara luas.',
         align='justify', indent=True)

    save(doc, 'BD 4 - LATAR BELAKANG INOVASI.docx')


# ══════════════════════════════════════════════════════════════
# 5. MEKANISME MONITORING DAN EVALUASI
# ══════════════════════════════════════════════════════════════
def buat_monev():
    doc = new_doc()
    doc_header(doc, 'BUKTI DUKUNG MEKANISME MONITORING DAN EVALUASI INOVASI')

    heading(doc, 'I. TUJUAN MONITORING DAN EVALUASI (MONEV)')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run('Monitoring dan evaluasi (Monev) ')
    r0.bold=True; r0.font.name='Times New Roman'; r0.font.size=Pt(12)
    r1 = p.add_run(
        'dilakukan untuk memastikan bahwa implementasi inovasi ANTRI-CENGKEH berjalan '
        'efektif, efisien, dan mencapai target yang telah ditetapkan. Monev juga '
        'berfungsi sebagai alat perbaikan berkelanjutan serta pengambilan kebijakan '
        'berbasis data antrian yang akurat.')
    r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'II. INSTRUMEN DAN MEKANISME MONEV')
    numbered(doc, 1, '', bold_prefix='Instrumen Monitoring:', sb=4, sa=2)
    for t in ['Dashboard antrian real-time sebagai pusat data antrian harian.',
              'Laporan CSV otomatis (jumlah antrian, jenis layanan, waktu tunggu rata-rata).',
              'Log aktivitas sistem yang mencatat seluruh transaksi antrian.']:
        sub_bullet(doc, t, level=2)

    numbered(doc, 2, '', bold_prefix='Mekanisme Monev:', sb=4, sa=2)

    monev_items = [
        ('Monitoring Harian:', [
            'Dilakukan oleh petugas PST melalui dashboard antrian setiap hari kerja.',
            'Meliputi pengecekan jumlah antrian per layanan dan ketertiban sistem.',
        ]),
        ('Monitoring Mingguan:', [
            'Admin sistem mengekspor dan menganalisis laporan CSV antrian.',
            'Laporan disampaikan kepada pimpinan OPP.',
        ]),
        ('Evaluasi Bulanan:', [
            'Dilakukan oleh pimpinan OPP dan tim pengelola sistem.',
            'Menganalisis efektivitas sistem, tren kunjungan, dan kepuasan pengunjung.',
        ]),
        ('Survei Kepuasan Internal:', [
            'Disebarkan kepada petugas dan pengguna internal setiap 3 bulan.',
            'Hasil digunakan untuk perbaikan sistem dan penyesuaian SOP.',
        ]),
    ]
    for title, subs in monev_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.5)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run('o  ' + title); r.bold=True
        r.font.name='Times New Roman'; r.font.size=Pt(12)
        for s in subs:
            pb = doc.add_paragraph()
            pb.paragraph_format.left_indent  = Cm(2.25)
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after  = Pt(1)
            rb = pb.add_run('▪  ' + s)
            rb.font.name='Times New Roman'; rb.font.size=Pt(12)

    add_hline(doc)
    heading(doc, 'III. INDIKATOR MONEV DAN PENGUKURAN KINERJA')
    add_table(doc,
        ['Indikator', 'Alat Ukur', 'Target'],
        [('Jumlah antrian tercatat',    'Dashboard / Laporan CSV',      '≥ 95% antrian tercatat'),
         ('Waktu rata-rata tunggu',     'Data timestamp sistem',         '< 15 menit per pengunjung'),
         ('Ketersediaan laporan harian','Sistem dashboard otomatis',     '100% tersedia tiap hari kerja'),
         ('Kepuasan petugas & pengguna','Survei internal (triwulan)',    '≥ 80% menyatakan puas'),
        ], col_widths=[5.5, 5.5, 5.0])

    add_hline(doc)
    heading(doc, 'IV. TINDAK LANJUT DAN PERBAIKAN BERKELANJUTAN')
    for t in ['Temuan dari monitoring dicatat dalam log sistem dan laporan petugas harian.',
              'Perbaikan fitur atau penyesuaian SOP dilakukan tiap bulan berdasarkan hasil evaluasi.',
              'Dokumentasi setiap update sistem disimpan bersama file sistem dan panduan penggunaan.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, 'V. PENUTUP', sb=6)
    para(doc,
         'Mekanisme monitoring dan evaluasi inovasi ANTRI-CENGKEH dilakukan secara '
         'sistematis, terintegrasi, dan berbasis data. Pendekatan ini menjamin '
         'keberlanjutan dan perbaikan berkelanjutan yang sangat penting dalam mendukung '
         'reformasi pelayanan publik BPS Kabupaten Tolitoli.',
         align='justify', indent=True)

    add_hline(doc)
    para(doc, 'Lampiran :', italic=True, align='left', sb=4, sa=2)
    for t in ['Template log monitoring harian antrian.',
              'Dashboard visual antrian (tangkapan layar).',
              'Contoh laporan CSV antrian harian.',
              'Rekap laporan mingguan/bulanan antrian PST.']:
        bullet(doc, t)

    save(doc, 'BD 5 - MEKANISME MONEV.docx')


# ══════════════════════════════════════════════════════════════
# 6. STRATEGI KEBERLANJUTAN
# ══════════════════════════════════════════════════════════════
def buat_keberlanjutan():
    doc = new_doc()
    para(doc, 'BUKTI DUKUNG STRATEGI KEBERLANJUTAN INOVASI', bold=True, align='left', sb=0, sa=8)
    para(doc, 'Judul Inovasi', bold=True, align='left', sb=0, sa=2)
    para(doc, 'ANTRI-CENGKEH', bold=True, align='left', sb=0, sa=2)
    para(doc, '(Antrian Terpadu Real-time Informatif — Cepat Efisien Nomor Gilirannya Kiosk Elektronik Harian)',
         align='left', sb=0, sa=6)
    add_hline(doc)

    heading(doc, '1. Latar Belakang')
    para(doc,
         'Inovasi ANTRI-CENGKEH hadir sebagai solusi antrian digital terpadu di lingkungan '
         'Pelayanan Statistik Terpadu (PST) BPS Kabupaten Tolitoli. Untuk menjaga '
         'keberlanjutan inovasi ini, strategi jangka pendek, menengah, dan panjang telah '
         'disusun guna memastikan inovasi berjalan konsisten dan memberi dampak jangka panjang.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, '2. Strategi Keberlanjutan Inovasi')
    heading(doc, 'A. Strategi Jangka Pendek (0–6 bulan):', sb=6)
    for t in ['Melakukan pelatihan internal kepada petugas layanan PST dan admin sistem ANTRI-CENGKEH.',
              'Menerapkan monitoring harian terhadap pelaksanaan dan ketertiban sistem antrian.',
              'Mengumpulkan masukan pengguna secara berkala untuk evaluasi awal dan perbaikan sistem.']:
        bullet(doc, t)

    heading(doc, 'B. Strategi Jangka Menengah (6–12 bulan):', sb=6)
    for t in ['Pengembangan lanjutan fitur sistem berdasarkan kebutuhan petugas dan pengunjung.',
              'Penyusunan dan sosialisasi SOP digitalisasi antrian dan evaluasi pelayanan harian.',
              'Penambahan indikator evaluasi untuk mengukur efektivitas inovasi secara berkala.']:
        bullet(doc, t)

    heading(doc, 'C. Strategi Jangka Panjang (>1 tahun):', sb=6)
    for t in ['Integrasi ANTRI-CENGKEH dengan sistem informasi layanan statistik BPS yang lebih luas.',
              'Replikasi inovasi ke BPS Kabupaten/Kota lainnya sebagai role model inovasi digital PST.',
              'Peningkatan keamanan data, cadangan data otomatis, dan skalabilitas sistem aplikasi.']:
        bullet(doc, t)

    add_hline(doc)
    heading(doc, '3. Penutup', sb=6)
    para(doc,
         'Dokumen ini menjadi bukti bahwa inovasi ANTRI-CENGKEH telah dirancang tidak hanya '
         'untuk menyelesaikan permasalahan antrian saat ini, tetapi juga untuk memberikan '
         'solusi berkelanjutan yang dapat diperluas, ditingkatkan, dan diterapkan secara luas '
         'di lingkungan kerja BPS maupun instansi pelayanan publik lainnya.',
         align='justify', indent=True)

    save(doc, 'BD 6 - STRATEGI KEBERLANJUTAN.docx')


# ══════════════════════════════════════════════════════════════
# 7. SUMBER DAYA INOVASI
# ══════════════════════════════════════════════════════════════
def buat_sumber_daya():
    doc = new_doc()
    para(doc, 'BUKTI DUKUNG SUMBER DAYA INOVASI', bold=True, size=14, align='left', sb=0, sa=6)
    mixed_para(doc, [('Judul Inovasi:\n', True, False)], align='left', sb=0, sa=2)
    p = doc.paragraphs[-1]
    r2 = p.add_run('ANTRI-CENGKEH ')
    r2.bold=True; r2.font.name='Times New Roman'; r2.font.size=Pt(12)
    r3 = p.add_run('(Antrian Terpadu Real-time Informatif — Cepat Efisien Nomor Gilirannya Kiosk Elektronik Harian)')
    r3.font.name='Times New Roman'; r3.font.size=Pt(12)
    add_hline(doc)

    heading(doc, '1. Ringkasan Inovasi')
    para(doc,
         'ANTRI-CENGKEH adalah sistem antrian digital terpadu untuk mempermudah pengelolaan '
         'antrian, pencatatan kunjungan, serta monitoring pelayanan di Pelayanan Statistik '
         'Terpadu (PST) BPS Kabupaten Tolitoli. Inovasi ini bertujuan untuk meningkatkan '
         'efisiensi pelayanan, ketertiban antrian, dan akuntabilitas pelaporan harian.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, '2. Sumber Daya yang Digunakan')
    add_table(doc,
        ['Jenis Sumber Daya', 'Deskripsi'],
        [(('SDM', True),
          'Tim pengembang terdiri dari pegawai IPDS, operator PST, dan koordinator inovasi BPS Kabupaten Tolitoli.'),
         (('Perangkat Lunak', True),
          'HTML5, CSS3, JavaScript ES6+ (gratis), Browser Chrome/Edge (gratis), Visual Studio Code (gratis).'),
         (('Perangkat Keras', True),
          'Laptop/PC kerja, tablet atau layar sentuh untuk kiosk antrian, monitor tambahan untuk layar antrian.'),
         (('Anggaran', True),
          'Non-anggaran (seluruh proses dilakukan dengan sumber daya internal, tanpa biaya tambahan).'),
        ], col_widths=[4.5, 11.5])

    add_hline(doc)
    heading(doc, '3. Bukti Pendukung')
    heading(doc, 'a. Foto/Video Dokumentasi', sb=4)
    para(doc,
         'Visualisasi proses kerja sistem ANTRI-CENGKEH yang digunakan di lingkungan kantor, '
         'termasuk demo kiosk pengambilan nomor, layar antrian, dashboard petugas, dan panel admin.',
         align='justify', indent=True)
    heading(doc, 'b. Dokumen Kerja dan Monitoring', sb=4)
    para(doc,
         'File sistem antrian digital (index.html, display.html, officer.html, admin.html, js/state.js), '
         'laporan CSV antrian harian, dan panduan penggunaan sistem.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, '4. Keterangan Peran Sumber Daya')
    add_table(doc,
        ['Nama / Jabatan', 'Peran'],
        [('Petugas Layanan PST',
          'Mengoperasikan dashboard antrian, memanggil dan mengelola antrian pengunjung PST.'),
         ('Tim IPDS',
          'Pengembang dan pemelihara sistem ANTRI-CENGKEH.'),
         ('Koordinator Fungsional',
          'Monitoring realisasi antrian dan evaluasi harian serta laporan mingguan kepada pimpinan.'),
        ], col_widths=[5.0, 11.0])

    save(doc, 'BD 7 - SUMBER DAYA INOVASI.docx')


# ══════════════════════════════════════════════════════════════
# 8. POTENSI REPLIKASI
# ══════════════════════════════════════════════════════════════
def buat_replikasi():
    doc = new_doc()
    para(doc, 'BUKTI DUKUNG POTENSI REPLIKASI INOVASI', bold=True, size=14, align='left', sb=0, sa=6)
    para(doc, 'Judul Inovasi:', bold=True, align='left', sb=0, sa=2)
    para(doc, 'ANTRI-CENGKEH', bold=True, align='left', sb=0, sa=2)
    para(doc,
         '(Antrian Terpadu Real-time Informatif — Cepat Efisien Nomor Gilirannya Kiosk Elektronik Harian)',
         align='left', sb=0, sa=6)
    add_hline(doc)

    heading(doc, '1. Deskripsi Singkat Inovasi')
    para(doc,
         'ANTRI-CENGKEH merupakan sistem antrian digital terpadu untuk pengelolaan antrian '
         'layanan di Pelayanan Statistik Terpadu (PST) BPS Kabupaten Tolitoli. Sistem ini '
         'dibangun menggunakan teknologi web (HTML5, CSS3, JavaScript) yang berjalan pada '
         'jaringan lokal kantor, memungkinkan efisiensi operasional tanpa memerlukan '
         'infrastruktur teknologi canggih maupun koneksi internet.',
         align='justify', indent=True)

    add_hline(doc)
    heading(doc, '2. Potensi Replikasi Inovasi')
    para(doc, 'Inovasi ini memiliki potensi replikasi yang tinggi karena:', align='left', sa=4)
    add_table(doc,
        ['Aspek Potensi', 'Penjelasan'],
        [(('Sederhana & Mudah Diterapkan', True),
          'Menggunakan teknologi web terbuka yang umum tersedia dan dapat berjalan di semua perangkat komputer modern tanpa instalasi khusus.'),
         (('Dokumentasi Lengkap', True),
          'Tersedia file sistem siap pakai, panduan penggunaan langkah demi langkah, dan SOP yang dapat langsung dibagikan.'),
         (('Skalabilitas Tinggi', True),
          'Dapat digunakan di seluruh unit kerja BPS Kabupaten/Kota di Indonesia, bahkan instansi pemerintah non-BPS.'),
         (('Fleksibel & Adaptif', True),
          'Sistem dapat dimodifikasi sesuai kebutuhan instansi (penambahan jenis layanan, pengaturan tampilan, dll).'),
         (('Tanpa Biaya Tambahan', True),
          'Tidak membutuhkan anggaran khusus karena memanfaatkan perangkat dan sumber daya internal yang sudah ada.'),
        ], col_widths=[5.0, 11.0])

    add_hline(doc)
    heading(doc, '3. Bukti Kesiapan Replikasi')
    kesiapan = [
        ('File Sistem Siap Pakai: ',
         'File ANTRI-CENGKEH (index.html, display.html, officer.html, admin.html, js/state.js) telah disusun dengan struktur siap pakai.'),
        ('Panduan Penggunaan: ',
         'Tersedia panduan langkah demi langkah untuk implementasi sistem di unit kerja lain.'),
        ('Video Tutorial: ',
         'Dokumentasi video penggunaan sistem untuk membantu proses sosialisasi dan replikasi.'),
        ('Komitmen Tim: ',
         'Tim inovator siap mendampingi proses replikasi di satuan kerja lain.'),
    ]
    for bold_t, norm_t in kesiapan:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.75)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r0 = p.add_run('•  ' + bold_t); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run(norm_t)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, '4. Rencana Replikasi')
    rencana = [
        ('Tahap 1 – Sosialisasi',
         'Penyebarluasan informasi dan manfaat inovasi melalui forum internal dan media sosial BPS.'),
        ('Tahap 2 – Adaptasi Lokal',
         'Modifikasi sistem berdasarkan karakteristik dan kebutuhan satuan kerja penerima.'),
        ('Tahap 3 – Pendampingan',
         'Tim inovator melakukan pendampingan implementasi awal dan pelatihan teknis kepada petugas.'),
        ('Tahap 4 – Evaluasi',
         'Monitoring bersama terhadap keberhasilan dan kendala dalam replikasi di unit kerja penerima.'),
    ]
    for i, (title, desc) in enumerate(rencana, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before      = Pt(3)
        p.paragraph_format.space_after       = Pt(3)
        r0 = p.add_run(f'{i}.  {title}\n'); r0.bold=True
        r0.font.name='Times New Roman'; r0.font.size=Pt(12)
        r1 = p.add_run('     ' + desc)
        r1.font.name='Times New Roman'; r1.font.size=Pt(12)

    add_hline(doc)
    heading(doc, '5. Penutup', sb=6)
    para(doc,
         'Dengan kesederhanaan sistem, ketersediaan dokumentasi, dan komitmen tim, inovasi '
         'ANTRI-CENGKEH sangat layak untuk direplikasi ke unit kerja lain guna meningkatkan '
         'layanan publik berbasis teknologi informasi.',
         align='justify', indent=True)

    save(doc, 'BD 8 - POTENSI REPLIKASI.docx')


# ══════════════════════════════════════════════════════════════
# RUN ALL
# ══════════════════════════════════════════════════════════════
print('Membuat 8 dokumen Bukti Dukung ANTRI-CENGKEH...\n')
buat_keunggulan()
buat_tujuan()
buat_cara_kerja()
buat_latar_belakang()
buat_monev()
buat_keberlanjutan()
buat_sumber_daya()
buat_replikasi()
print('\nSelesai! Semua 8 file berhasil dibuat.')
