import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r'C:\Users\ERIKA A K\Documents\7206\Sistem Antrian'
NAVY  = '#003d82'
BLUE  = '#1565c0'
GREEN = '#2e7d32'
GOLD  = '#f9a825'
RED   = '#c62828'
GREY  = '#eef2f7'

# ══════════════════════════════════════════════════════════════
# DIAGRAM HELPERS
# ══════════════════════════════════════════════════════════════

def savefig(fig, name):
    path = os.path.join(BASE, name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def rbox(ax, cx, cy, w, h, fc, ec, text, fs=9, tc='white', bold=True, lw=1.5):
    r = FancyBboxPatch((cx-w/2, cy-h/2), w, h,
                       boxstyle="round,pad=0.08",
                       facecolor=fc, edgecolor=ec, linewidth=lw, zorder=3)
    ax.add_patch(r)
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=fs, color=tc, fontweight='bold' if bold else 'normal',
            multialignment='center', zorder=4)

def arr(ax, x1, y1, x2, y2, color=NAVY, lw=1.8, style='->', ls='-'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color,
                                lw=lw, linestyle=ls),
                zorder=2)

def stick_actor(ax, cx, cy, label, color):
    """Simple stick figure actor."""
    # head
    head = plt.Circle((cx, cy+0.42), 0.18, color=color, zorder=5)
    ax.add_patch(head)
    # body
    ax.plot([cx, cx],      [cy+0.24, cy-0.25], color=color, lw=2, zorder=5)
    # arms
    ax.plot([cx-0.28, cx+0.28], [cy+0.08, cy+0.08], color=color, lw=2, zorder=5)
    # legs
    ax.plot([cx, cx-0.22], [cy-0.25, cy-0.55], color=color, lw=2, zorder=5)
    ax.plot([cx, cx+0.22], [cy-0.25, cy-0.55], color=color, lw=2, zorder=5)
    ax.text(cx, cy-0.78, label, ha='center', va='top',
            fontsize=8.5, fontweight='bold', color=color, zorder=5)


# ══════════════════════════════════════════════════════════════
# DIAGRAM 1: ARSITEKTUR
# ══════════════════════════════════════════════════════════════
def diag_arsitektur():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6.5); ax.axis('off')
    fig.patch.set_facecolor('white')

    # ── Pengunjung ──
    stick_actor(ax, 1.0, 4.5, 'Pengunjung', BLUE)
    arr(ax, 1.55, 4.7, 2.5, 4.7)
    ax.text(2.0, 4.88, 'Pilih layanan\n& ambil nomor',
            ha='center', fontsize=7.5, color='gray')

    # ── Kiosk ──
    rbox(ax, 3.5, 4.7, 2.0, 0.75, BLUE, BLUE,
         'Kiosk Digital\nindex.html', fs=9)

    # ── arrow kiosk → localStorage ──
    arr(ax, 3.5, 4.32, 3.5, 3.55)
    ax.text(3.65, 3.9, 'addQueue()', ha='left', fontsize=7.5, color='gray')

    # ── localStorage ──
    rbox(ax, 5.5, 3.2, 2.4, 0.75, GOLD, GOLD,
         'localStorage\nbps_tolitoli_antrian', fs=9, tc='#1a2332')
    arr(ax, 3.5, 3.2, 4.28, 3.2)   # kiosk → storage (horizontal)

    # ── BroadcastChannel note ──
    ax.text(5.5, 2.65, '← BroadcastChannel API (sinkronisasi real-time) →',
            ha='center', fontsize=7.5, color='gray', style='italic')

    # ── Three output boxes ──
    rbox(ax, 1.8, 1.4, 2.2, 0.7, '#0277bd', '#0277bd',
         'Layar Antrian\ndisplay.html', fs=9)
    rbox(ax, 5.5, 1.4, 2.2, 0.7, GREEN, GREEN,
         'Dashboard Petugas\nofficer.html', fs=9)
    rbox(ax, 9.2, 1.4, 2.2, 0.7, RED, RED,
         'Panel Admin\nadmin.html', fs=9)

    # arrows localStorage → 3 boxes
    arr(ax, 4.3, 2.95, 2.5, 1.78)
    arr(ax, 5.5, 2.82, 5.5, 1.77)
    arr(ax, 6.7, 2.95, 8.5, 1.78)

    ax.text(3.1, 2.28, 'Display\nNomor', ha='center', fontsize=7.5, color='gray')
    ax.text(5.65, 2.28, 'Kelola\nAntrian', ha='left', fontsize=7.5, color='gray')
    ax.text(8.0, 2.28, 'Laporan &\nMonitor', ha='center', fontsize=7.5, color='gray')

    # ── Petugas actor ──
    stick_actor(ax, 5.5, 0.0, 'Petugas', GREEN)
    arr(ax, 5.5, 0.58, 5.5, 1.04)

    # ── Admin actor ──
    stick_actor(ax, 9.2, 0.0, 'Admin', RED)
    arr(ax, 9.2, 0.58, 9.2, 1.04)

    # Title
    ax.text(5.5, 6.25, 'Gambar 1. Rancangan Sistem ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    return savefig(fig, '_diag1_arsitektur.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 2: SPESIFIKASI TEKNOLOGI
# ══════════════════════════════════════════════════════════════
def diag_teknologi():
    fig, ax = plt.subplots(figsize=(9, 5.5))
    ax.set_xlim(0, 9); ax.set_ylim(0, 5.5); ax.axis('off')
    fig.patch.set_facecolor('white')

    # Header bar
    rbox(ax, 4.5, 5.1, 8.6, 0.65, NAVY, NAVY,
         'Spesifikasi Teknologi ANTRI-CENGKEH', fs=12, tc='white')

    # Background card
    bg = FancyBboxPatch((0.15, 0.2), 8.7, 4.55,
                        boxstyle="round,pad=0.1",
                        facecolor=GREY, edgecolor=NAVY, linewidth=1.5)
    ax.add_patch(bg)

    items = [
        (NAVY,  '[HW]', 'Perangkat Keras',
         '• PC/Laptop (Petugas & Admin)\n• Tablet/Layar Sentuh (Kiosk)\n• Monitor tambahan (Display Antrian)'),
        (BLUE,  '[CODE]', 'Bahasa & Teknologi',
         '• HTML5, CSS3, JavaScript ES6+\n• localStorage API\n• BroadcastChannel API\n• Web Speech API (pengumuman suara)'),
        (GREEN, '[NET]', 'Jaringan',
         '• LAN / WiFi lokal kantor\n• Tidak memerlukan koneksi internet\n• Cukup 1 jaringan lokal yang sama'),
        (RED,   '[WEB]', 'Browser yang Didukung',
         '• Google Chrome (direkomendasikan)\n• Microsoft Edge\n• Mozilla Firefox\n• Versi terbaru / modern'),
    ]

    positions = [(1.35, 3.6), (5.15, 3.6), (1.35, 1.55), (5.15, 1.55)]

    for (color, icon, title, desc), (cx, cy) in zip(items, positions):
        # Card
        card = FancyBboxPatch((cx-1.8, cy-1.2), 3.6, 2.4,
                              boxstyle="round,pad=0.1",
                              facecolor='white', edgecolor=color, linewidth=2)
        ax.add_patch(card)
        # Title bar inside card
        tbar = FancyBboxPatch((cx-1.8, cy+0.8), 3.6, 0.4,
                              boxstyle="round,pad=0.0",
                              facecolor=color, edgecolor=color)
        ax.add_patch(tbar)
        ax.text(cx, cy+1.0, f'{icon}  {title}',
                ha='center', va='center', fontsize=9.5,
                fontweight='bold', color='white')
        ax.text(cx, cy-0.1, desc,
                ha='center', va='center', fontsize=8.5, color='#333',
                multialignment='left')

    ax.text(4.5, 5.42, 'Gambar 2. Spesifikasi Teknologi ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    return savefig(fig, '_diag2_teknologi.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 3: USE CASE
# ══════════════════════════════════════════════════════════════
def diag_usecase():
    fig, ax = plt.subplots(figsize=(13, 9))
    ax.set_xlim(0, 13); ax.set_ylim(0, 9.5); ax.axis('off')
    fig.patch.set_facecolor('white')

    # System boundary
    sb = FancyBboxPatch((2.1, 0.8), 10.5, 7.8,
                        boxstyle="round,pad=0.1",
                        facecolor='#f8faff', edgecolor=NAVY, linewidth=2,
                        linestyle='--', zorder=1)
    ax.add_patch(sb)
    ax.text(7.35, 8.85, '« system »  ANTRI-CENGKEH',
            ha='center', fontsize=10, color=NAVY, style='italic')

    # ── PENGUNJUNG use cases ──
    peng_uc = [
        (4.2, 7.5, 'Pilih Jenis Layanan'),
        (4.2, 6.6, 'Ambil Nomor Antrian'),
        (4.2, 5.7, 'Lihat Tiket Digital'),
        (4.2, 4.8, 'Lihat Layar Antrian'),
    ]
    for x, y, t in peng_uc:
        rbox(ax, x, y, 2.6, 0.55, 'white', BLUE, t, fs=8.5, tc=BLUE, bold=False)

    # ── PETUGAS use cases ──
    pet_uc = [
        (7.35, 7.5, 'Input Identitas Petugas'),
        (7.35, 6.6, 'Panggil Antrian Berikutnya'),
        (7.35, 5.7, 'Panggil Antrian Spesifik'),
        (7.35, 4.8, 'Selesaikan Layanan'),
        (7.35, 3.9, 'Lewati / Skip Antrian'),
        (7.35, 3.0, 'Panggil Ulang Antrian'),
    ]
    for x, y, t in pet_uc:
        rbox(ax, x, y, 2.8, 0.55, 'white', GREEN, t, fs=8.5, tc=GREEN, bold=False)

    # ── ADMIN use cases ──
    adm_uc = [
        (10.8, 7.5, 'Login Admin'),
        (10.8, 6.6, 'Lihat Dashboard'),
        (10.8, 5.7, 'Lihat Laporan Antrian'),
        (10.8, 4.8, 'Export CSV Antrian'),
        (10.8, 3.9, 'Reset Data Antrian'),
        (10.8, 3.0, 'Kelola Akun Admin'),
    ]
    for x, y, t in adm_uc:
        rbox(ax, x, y, 2.6, 0.55, 'white', RED, t, fs=8.5, tc=RED, bold=False)

    # ── Actors ──
    stick_actor(ax, 0.85, 5.5, 'Pengunjung', BLUE)
    stick_actor(ax, 7.35, 0.0, 'Petugas', GREEN)   # below
    stick_actor(ax, 12.3, 5.0, 'Admin', RED)

    # Lines: Pengunjung → use cases
    for x, y, t in peng_uc:
        ax.plot([1.3, x-1.3], [5.7, y], color=BLUE, lw=1, alpha=0.5, zorder=2)

    # Lines: Petugas → use cases
    for x, y, t in pet_uc:
        ax.plot([7.35, x], [0.65, y-0.28], color=GREEN, lw=1, alpha=0.5, zorder=2)

    # Lines: Admin → use cases
    for x, y, t in adm_uc:
        ax.plot([12.1, x+1.3], [5.2, y], color=RED, lw=1, alpha=0.5, zorder=2)

    # Column labels
    ax.text(4.2,  8.5, 'Pengunjung', ha='center', fontsize=10,
            fontweight='bold', color=BLUE)
    ax.text(7.35, 8.5, 'Petugas', ha='center', fontsize=10,
            fontweight='bold', color=GREEN)
    ax.text(10.8, 8.5, 'Admin', ha='center', fontsize=10,
            fontweight='bold', color=RED)

    ax.text(6.5, 9.25, 'Gambar 3. Diagram Use Case ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    return savefig(fig, '_diag3_usecase.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 4: PERANCANGAN DATA
# ══════════════════════════════════════════════════════════════
def diag_data():
    fig, ax = plt.subplots(figsize=(12, 7.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7.5); ax.axis('off')
    fig.patch.set_facecolor('white')

    ax.text(6, 7.2, 'Gambar 4. Rancangan Data localStorage ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    # ── State Object (kiri) ──
    state_fields = [
        ('queues',          'Array',   'Daftar semua objek antrian'),
        ('counters',        'Object',  '{ K:0, P:0, R:0, L:0 }'),
        ('currentServing',  'Object',  'Antrian yg sedang dilayani'),
        ('officer',         'Object',  '{ name: "", email: "" }'),
        ('lastUpdated',     'number',  'Timestamp update terakhir (ms)'),
    ]

    queue_fields = [
        ('id',            'string',  'Contoh: K-001, P-003'),
        ('code',          'string',  'K / P / R / L'),
        ('number',        'number',  'Nomor urut per kode'),
        ('typeName',      'string',  'Nama jenis layanan'),
        ('desk',          'string',  'Nama meja pelayanan'),
        ('name',          'string',  'Nama pengunjung'),
        ('phone',         'string',  'Nomor HP pengunjung'),
        ('institution',   'string',  'Instansi / asal daerah'),
        ('gender',        'string',  'Laki-laki / Perempuan'),
        ('purpose',       'string',  'Keperluan kunjungan'),
        ('status',        'string',  'waiting/serving/done/skip'),
        ('createdAt',     'number',  'Timestamp ambil nomor'),
        ('calledAt',      'number',  'Timestamp saat dipanggil'),
        ('doneAt',        'number',  'Timestamp selesai layanan'),
        ('waitingBefore', 'number',  'Jumlah antrian sebelumnya'),
    ]

    def draw_table(ax, lx, ty, title, fields, color):
        row_h = 0.37
        header_h = 0.5
        body_h = len(fields) * row_h
        total_h = header_h + body_h
        w = 5.5

        # Header
        hdr = FancyBboxPatch((lx, ty - header_h), w, header_h,
                             boxstyle="round,pad=0.05",
                             facecolor=color, edgecolor=color, linewidth=0, zorder=3)
        ax.add_patch(hdr)
        ax.text(lx + w/2, ty - header_h/2, title,
                ha='center', va='center', fontsize=9.5,
                fontweight='bold', color='white', zorder=4)

        # Body background
        body = FancyBboxPatch((lx, ty - total_h), w, body_h,
                              boxstyle="round,pad=0.0",
                              facecolor='white', edgecolor=color, linewidth=1.5, zorder=3)
        ax.add_patch(body)

        for i, (fname, ftype, fdesc) in enumerate(fields):
            fy = ty - header_h - (i + 0.5) * row_h
            # alternating row bg
            if i % 2 == 0:
                row_bg = FancyBboxPatch((lx, ty - header_h - (i+1)*row_h),
                                       w, row_h, boxstyle="round,pad=0.0",
                                       facecolor='#f8faff', edgecolor='none', zorder=3)
                ax.add_patch(row_bg)
            ax.text(lx + 0.12, fy, fname, ha='left', va='center',
                    fontsize=8, color=color, fontweight='bold', zorder=4)
            ax.text(lx + 1.5,  fy, ftype, ha='left', va='center',
                    fontsize=7.5, color='gray', zorder=4)
            ax.text(lx + 2.4,  fy, fdesc, ha='left', va='center',
                    fontsize=7.5, color='#333', zorder=4)

        return total_h

    h1 = draw_table(ax, 0.2,  6.8, 'State Utama  (key: bps_tolitoli_antrian)',
                   state_fields, NAVY)
    h2 = draw_table(ax, 6.3,  6.8, 'Objek Antrian  (queues[ ])',
                   queue_fields, BLUE)

    # Arrow: queues field → Queue object
    arrow_y = 6.8 - 0.5 - 0.37*0.5   # row 0 (queues) center y
    arr(ax, 5.7, arrow_y, 6.28, arrow_y, color=GOLD, lw=2)
    ax.text(6.0, arrow_y + 0.12, 'tiap\nelemen', ha='center',
            fontsize=7.5, color='gray')

    return savefig(fig, '_diag4_data.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 5: STRUKTUR FILE / KODE
# ══════════════════════════════════════════════════════════════
def diag_kode():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6); ax.axis('off')
    fig.patch.set_facecolor('white')

    # Title bar
    rbox(ax, 5.5, 5.7, 10.6, 0.5, NAVY, NAVY,
         'Tools & Struktur File Proyek ANTRI-CENGKEH', fs=11)

    # Left: IDE & tools
    rbox(ax, 2.2, 4.85, 4.0, 0.55, BLUE, BLUE, 'IDE: Visual Studio Code', fs=10)
    tools = [
        ('[LANG]', 'Bahasa Pemrograman',
         'HTML5, CSS3, JavaScript (ES6+)'),
        ('[DB]', 'Penyimpanan Data',
         'Web localStorage (tanpa database server)'),
        ('[SYNC]', 'Sinkronisasi Real-time',
         'BroadcastChannel API'),
        ('[AUDIO]', 'Pengumuman Suara',
         'Web Speech Synthesis API'),
    ]
    for i, (icon, label, val) in enumerate(tools):
        y = 4.1 - i * 0.75
        rbox(ax, 2.2, y, 4.0, 0.55, GREY, BLUE,
             f'{icon}  {label}\n{val}', fs=8.5, tc=NAVY, bold=False)

    # Right: file tree
    rbox(ax, 8.0, 4.85, 5.0, 0.55, GREEN, GREEN, 'Struktur File Proyek', fs=10)

    tree = [
        ('[DIR] Sistem Antrian/', NAVY,   True),
        ('   +-- index.html',   BLUE,  False),
        ('   |     Kiosk -- pengambilan nomor antrian', '#555', False),
        ('   +-- display.html', BLUE,  False),
        ('   |     Layar antrian ruang tunggu', '#555', False),
        ('   +-- officer.html', GREEN, False),
        ('   |     Dashboard petugas pelayanan', '#555', False),
        ('   +-- admin.html',   RED,   False),
        ('   |     Panel administrator sistem', '#555', False),
        ('   +-- [DIR] js/', NAVY,  False),
        ('        +-- state.js   (shared state manager)', '#555', False),
    ]
    for i, (txt, color, is_bold) in enumerate(tree):
        y = 4.22 - i * 0.37
        ax.text(5.6, y, txt, ha='left', va='center',
                fontsize=8.2, color=color,
                fontweight='bold' if is_bold else 'normal',
                fontfamily='monospace')

    ax.text(5.5, 0.25, 'Gambar 5. Tools Pengembangan & Struktur File ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    return savefig(fig, '_diag5_kode.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 6: ALUR PENGGUNAAN (pengganti "Upload ke Playstore")
# ══════════════════════════════════════════════════════════════
def diag_alur():
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.set_xlim(0, 11); ax.set_ylim(0, 5); ax.axis('off')
    fig.patch.set_facecolor('white')

    rbox(ax, 5.5, 4.7, 10.6, 0.55, NAVY, NAVY,
         'Cara Penggunaan / Deployment ANTRI-CENGKEH', fs=11)

    steps = [
        (BLUE,  '1', 'Persiapan\nJaringan',
         '• Hubungkan semua\n  perangkat ke WiFi/LAN\n  yang sama'),
        (BLUE,  '2', 'Buka Kiosk\n(index.html)',
         '• Buka di browser tablet\n• Pengunjung pilih layanan\n  & ambil nomor'),
        (GREEN, '3', 'Buka Layar\n(display.html)',
         '• Tampilkan di monitor\n  ruang tunggu\n• Update otomatis'),
        (GREEN, '4', 'Buka Dashboard\n(officer.html)',
         '• Buka di PC petugas\n• Panggil & kelola\n  antrian'),
        (RED,   '5', 'Buka Admin\n(admin.html)',
         '• Login admin\n• Monitor & ekspor\n  laporan CSV'),
    ]

    for i, (color, num, title, desc) in enumerate(steps):
        cx = 1.05 + i * 2.1
        # Circle number
        circle = plt.Circle((cx, 3.65), 0.32, color=color, zorder=4)
        ax.add_patch(circle)
        ax.text(cx, 3.65, num, ha='center', va='center',
                fontsize=14, fontweight='bold', color='white', zorder=5)
        # Title
        ax.text(cx, 3.1, title, ha='center', va='top',
                fontsize=8.5, fontweight='bold', color=color,
                multialignment='center')
        # Desc box
        db = FancyBboxPatch((cx-0.95, 0.4), 1.9, 1.85,
                            boxstyle="round,pad=0.08",
                            facecolor='white', edgecolor=color, linewidth=1.5, zorder=3)
        ax.add_patch(db)
        ax.text(cx, 1.32, desc, ha='center', va='center',
                fontsize=8, color='#333', multialignment='left', zorder=4)

        if i < 4:
            arr(ax, cx+0.95, 3.65, cx+1.15, 3.65, color='#aaa')

    ax.text(5.5, 0.1, 'Gambar 6. Alur Penggunaan / Cara Menjalankan Sistem ANTRI-CENGKEH',
            ha='center', fontsize=11, fontweight='bold', color=NAVY)

    return savefig(fig, '_diag6_alur.png')


# ══════════════════════════════════════════════════════════════
# DIAGRAM 7: IMPLEMENTASI — Tampilan Sistem
# ══════════════════════════════════════════════════════════════
def diag_implementasi():
    fig, axes = plt.subplots(1, 4, figsize=(14, 5))
    fig.patch.set_facecolor('white')
    fig.suptitle('Gambar 7. Tampilan Sistem ANTRI-CENGKEH',
                 fontsize=12, fontweight='bold', color=NAVY, y=0.02)

    pages = [
        (BLUE,  'index.html\nKiosk Antrian',
         ['• Pilih Jenis Layanan', '• PST / Lainnya', '• Konsultasi Statistik',
          '• Perpustakaan Statistik', '• Rekomendasi Statistik', '',
          '• Langsung cetak tiket', '  dengan nomor antrian']),
        ('#0277bd', 'display.html\nLayar Antrian',
         ['• Tampil nomor aktif', '• Ukuran besar & jelas', '• Update real-time',
          '• Statistik per layanan', '', '• Daftar antrian tunggu',
          '• Jam & tanggal live', '• Running text info']),
        (GREEN, 'officer.html\nDashboard Petugas',
         ['• Identitas petugas', '• Tombol Panggil Berikutnya', '• Per jenis layanan',
          '', '• Tabel daftar antrian', '• Filter per tab', '',
          '• Selesai / Skip / Ulang']),
        (RED,   'admin.html\nPanel Admin',
         ['• Login aman', '• Dashboard statistik', '• Grafik per layanan',
          '', '• Laporan lengkap', '• Filter & pencarian',
          '• Export CSV', '• Reset & kelola akun']),
    ]

    for ax, (color, title, items) in zip(axes, pages):
        ax.set_facecolor(GREY)
        # Header
        hdr = FancyBboxPatch((0.02, 0.78), 0.96, 0.2,
                             boxstyle="round,pad=0.01",
                             facecolor=color, edgecolor=color,
                             transform=ax.transAxes, zorder=3)
        ax.add_patch(hdr)
        ax.text(0.5, 0.885, title, ha='center', va='center',
                transform=ax.transAxes, fontsize=8.5,
                fontweight='bold', color='white', zorder=4,
                multialignment='center')
        # Content
        for i, item in enumerate(items):
            y = 0.73 - i * 0.085
            ax.text(0.06, y, item, ha='left', va='top',
                    transform=ax.transAxes, fontsize=7.5,
                    color=NAVY if item.startswith('•') else '#666')
        ax.set_xlim(0, 1); ax.set_ylim(0, 1)
        ax.axis('off')
        for spine in ax.spines.values():
            spine.set_edgecolor(color)
            spine.set_linewidth(2)
            spine.set_visible(True)

    plt.tight_layout(rect=[0, 0.08, 1, 1])
    return savefig(fig, '_diag7_implementasi.png')


# ══════════════════════════════════════════════════════════════
# RUN ALL DIAGRAMS
# ══════════════════════════════════════════════════════════════
print('Membuat diagram 1/7: Arsitektur...')
p1 = diag_arsitektur()
print('Membuat diagram 2/7: Teknologi...')
p2 = diag_teknologi()
print('Membuat diagram 3/7: Use Case...')
p3 = diag_usecase()
print('Membuat diagram 4/7: Data...')
p4 = diag_data()
print('Membuat diagram 5/7: Kode...')
p5 = diag_kode()
print('Membuat diagram 6/7: Alur...')
p6 = diag_alur()
print('Membuat diagram 7/7: Implementasi...')
p7 = diag_implementasi()
print('Semua diagram selesai.\n')


# ══════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ══════════════════════════════════════════════════════════════
print('Membuat dokumen Word...')
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(4.0)
    section.right_margin  = Cm(3.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def p(doc, text='', align='justify', bold=False, fs=12, space_before=0, space_after=6, indent=True):
    para = doc.add_paragraph()
    if align == 'center':  para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify': para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if indent and align == 'justify':
        para.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = para.add_run(text)
        r.bold = bold
        r.font.name = 'Times New Roman'
        r.font.size = Pt(fs)
    return para

def h(doc, text, level=1, space_before=12, space_after=6):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    r = para.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return para

def img(doc, path, caption, width_cm=14):
    run = doc.add_paragraph().add_run()
    run.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for r in cap.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

def bullet(doc, text, indent_cm=1.0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent   = Cm(indent_cm)
    para.paragraph_format.space_before  = Pt(0)
    para.paragraph_format.space_after   = Pt(2)
    r = para.add_run(f'• {text}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)


# ─── JUDUL ───
p(doc, 'STUDI RANCANG BANGUN / BLUE PRINT', align='center', bold=True, fs=14,
  space_before=0, space_after=2, indent=False)
p(doc, 'SISTEM ANTRI-CENGKEH', align='center', bold=True, fs=14,
  space_before=0, space_after=16, indent=False)

# ─── Header nama inovasi ───
ph = doc.add_paragraph()
ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
ph.paragraph_format.space_before = Pt(0)
ph.paragraph_format.space_after  = Pt(8)
rr = ph.add_run(
    'SISTEM ANTRIAN DIGITAL PELAYANAN STATISTIK TERPADU '
    '(ANTRI-CENGKEH)\nBPS KABUPATEN TOLITOLI')
rr.bold = True
rr.font.name = 'Times New Roman'
rr.font.size = Pt(12)

# ─── Paragraf pengantar ───
p(doc,
  'Sistem Antrian Digital Pelayanan Statistik Terpadu (ANTRI-CENGKEH) adalah inovasi '
  'yang dibangun oleh BPS Kabupaten Tolitoli untuk meningkatkan kualitas dan efisiensi '
  'layanan antrian di Pelayanan Statistik Terpadu (PST). Inovasi ini hadir sebagai solusi '
  'atas permasalahan antrian manual yang selama ini dilakukan menggunakan daftar tunggu '
  'kertas sehingga tidak efisien, tidak terstruktur, dan sulit dievaluasi.')

p(doc,
  'ANTRI-CENGKEH dibangun berbasis teknologi web (HTML5, CSS3, JavaScript) yang berjalan '
  'pada jaringan lokal kantor tanpa memerlukan koneksi internet maupun biaya server. '
  'Nama CENGKEH terinspirasi dari identitas daerah Kabupaten Tolitoli sebagai Kota '
  'Cengkeh, sekaligus merupakan akronim dari: Cepat, Efisien, Nomor Gilirannya, Kiosk, '
  'Elektronik, Harian. Sistem ini terdiri dari empat modul utama yaitu Kiosk '
  'Pengambilan Nomor, Layar Antrian Digital, Dashboard Petugas, dan Panel Administrator.')

p(doc,
  'Dalam pembuatan sistem ini, tahapan kegiatan yang dilakukan dirinci sebagai berikut:',
  indent=False, space_after=8)

# ─── 1. PERANCANGAN ───
h(doc, '1.  Perancangan Sistem')
p(doc,
  'Perancangan sistem merupakan diagram alur bagaimana cara kerja ANTRI-CENGKEH '
  'dari pengunjung mengambil nomor hingga petugas memanggil dan menyelesaikan '
  'layanan. Selain itu, rancangan sistem membantu pengguna untuk memahami modul '
  'apa saja yang tersedia dan bagaimana data mengalir antar modul.')
p(doc,
  'Tampak pada Gambar 1 bahwa sistem ANTRI-CENGKEH memiliki dua kelompok pengguna '
  'utama, yaitu Pengunjung dan Petugas/Admin. Pengunjung mengambil nomor melalui '
  'Kiosk Digital, data antrian disimpan di localStorage, dan secara otomatis '
  'disinkronkan ke Layar Antrian, Dashboard Petugas, serta Panel Admin melalui '
  'BroadcastChannel API sehingga semua halaman selalu menampilkan data terkini '
  'tanpa perlu di-refresh secara manual.')
img(doc, p1, 'Gambar 1. Rancangan Sistem ANTRI-CENGKEH')

# ─── 2. SPESIFIKASI TEKNOLOGI ───
h(doc, '2.  Penyusunan Spesifikasi Teknologi yang Digunakan')
p(doc,
  'Spesifikasi teknologi yang digunakan merupakan spesifikasi minimal yang harus '
  'disiapkan sebagai syarat menjalankan sistem ANTRI-CENGKEH. Sistem dirancang '
  'seringan mungkin sehingga dapat berjalan pada perangkat komputer yang sudah '
  'tersedia di kantor tanpa memerlukan pembelian perangkat lunak berbayar. '
  'Spesifikasi teknologi yang digunakan dapat dilihat pada Gambar 2:')
img(doc, p2, 'Gambar 2. Spesifikasi Teknologi ANTRI-CENGKEH')

# ─── 3. USE CASE ───
h(doc, '3.  Use Case ANTRI-CENGKEH')
p(doc,
  'Use Case Diagram menggambarkan fungsionalitas sebuah sistem dari sudut pandang '
  'pengguna. Use Case Diagram sangat membantu ketika menyusun kebutuhan sistem, '
  'mengkomunikasikan rancangan dengan pengguna, dan merancang fitur yang ada pada '
  'sistem. Berikut Use Case Diagram Sistem ANTRI-CENGKEH:')
img(doc, p3, 'Gambar 3. Diagram Use Case ANTRI-CENGKEH', width_cm=15)
p(doc,
  'Terdapat tiga peran pengguna pada sistem ANTRI-CENGKEH. Penjelasan tentang '
  'hal-hal yang dapat dilakukan oleh masing-masing pengguna adalah sebagai berikut:')
bullet(doc, 'Pengunjung dapat memilih jenis layanan, mengambil nomor antrian secara '
            'mandiri melalui kiosk, melihat tiket digital, dan memantau layar antrian.')
bullet(doc, 'Petugas dapat menginput identitas, memanggil antrian berikutnya per jenis '
            'layanan, memanggil antrian spesifik, menyelesaikan, melewati, atau memanggil '
            'ulang antrian yang sedang dilayani.')
bullet(doc, 'Admin dapat login ke panel khusus, memantau dashboard statistik antrian, '
            'melihat laporan lengkap, mengekspor data CSV, mereset data harian, '
            'serta mengelola akun administrator.')

# ─── 4. PERANCANGAN DATA ───
h(doc, '4.  Perancangan Data')
p(doc,
  'Data pada sistem ANTRI-CENGKEH disimpan menggunakan Web localStorage yang '
  'tersedia secara bawaan pada seluruh browser modern. Pendekatan ini dipilih '
  'karena tidak memerlukan instalasi database server, sehingga sistem dapat '
  'berjalan sepenuhnya pada jaringan lokal tanpa infrastruktur tambahan. '
  'Semua halaman (Kiosk, Layar Antrian, Dashboard Petugas, dan Panel Admin) '
  'mengakses sumber data yang sama melalui modul state.js yang terpusat. '
  'Rancangan data dapat dilihat pada Gambar 4:')
img(doc, p4, 'Gambar 4. Rancangan Data localStorage ANTRI-CENGKEH', width_cm=15)

# ─── 5. PENULISAN KODE ───
h(doc, '5.  Penulisan Kode Program')
p(doc,
  'Penulisan kode program menggunakan Visual Studio Code sebagai Integrated '
  'Development Environment (IDE) yang digunakan sebagai alat bantu dalam '
  'melakukan coding. Sistem menggunakan bahasa pemrograman HTML5, CSS3, dan '
  'JavaScript ES6+ tanpa framework atau library eksternal sehingga tidak '
  'memerlukan proses instalasi atau build. Struktur file proyek dan tools '
  'yang digunakan dapat dilihat pada Gambar 5:')
img(doc, p5, 'Gambar 5. Tools Pengembangan & Struktur File ANTRI-CENGKEH')

# ─── 6. CARA PENGGUNAAN ───
h(doc, '6.  Cara Penggunaan / Deployment Sistem')
p(doc,
  'Sistem ANTRI-CENGKEH tidak memerlukan proses instalasi atau upload ke server '
  'publik. Cukup dengan membuka file HTML di browser pada perangkat yang terhubung '
  'ke jaringan lokal yang sama, sistem langsung siap digunakan. Hal ini menjadikan '
  'ANTRI-CENGKEH sangat mudah dan cepat untuk dioperasikan setiap hari kerja '
  'oleh seluruh petugas PST BPS Kabupaten Tolitoli. Alur penggunaan sistem '
  'dapat dilihat pada Gambar 6:')
img(doc, p6, 'Gambar 6. Alur Penggunaan / Cara Menjalankan ANTRI-CENGKEH')

# ─── 7. IMPLEMENTASI ───
h(doc, '7.  Implementasi Sistem')
p(doc,
  'Setelah sistem selesai dibuat, tahap berikutnya adalah implementasi sistem '
  'secara langsung di lingkungan PST BPS Kabupaten Tolitoli. Berikut adalah '
  'tampilan dari masing-masing modul sistem ANTRI-CENGKEH yang telah dibuat:')
img(doc, p7, 'Gambar 7. Tampilan Sistem ANTRI-CENGKEH', width_cm=15)
p(doc,
  'Keempat modul sistem berjalan secara bersamaan dan sinkron. Pengunjung '
  'mengambil nomor di Kiosk Antrian, nomor tampil real-time di Layar Antrian, '
  'petugas memanggil dan mengelola dari Dashboard Petugas, dan admin memantau '
  'serta mengunduh laporan dari Panel Admin. Seluruh proses berjalan otomatis '
  'tanpa memerlukan refresh halaman manual.')

# ─── SAVE ───
out = os.path.join(BASE, 'Studi Rancang Bangun ANTRI-CENGKEH.docx')
doc.save(out)
print(f'Dokumen berhasil disimpan: {out}')

# Clean up temp diagram files
for f in [p1, p2, p3, p4, p5, p6, p7]:
    try: os.remove(f)
    except: pass
print('Selesai.')
